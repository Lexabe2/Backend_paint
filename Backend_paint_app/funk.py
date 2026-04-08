from .models import ATM, Request, StatusReq, StatusATM
from datetime import date
from django.core.exceptions import ObjectDoesNotExist
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openpyxl.styles import Alignment, Border, Side
from openpyxl.utils import get_column_letter
import re

months = {
    1: "январь", 2: "февраль", 3: "март", 4: "апрель",
    5: "май", 6: "июнь", 7: "июль", 8: "август",
    9: "сентябрь", 10: "октябрь", 11: "ноябрь", 12: "декабрь"
}


def changes_req(req, status, request):
    Request.objects.filter(request_id=req).update(status=status)
    ATM.objects.filter(request=Request.objects.get(request_id=req)).update(status=status)
    StatusReq.objects.create(status=status, date_change=date.today(), user=request.user,
                             request=Request.objects.get(request_id=req))
    atms = ATM.objects.filter(request=Request.objects.get(request_id=req))
    for atm in atms:
        StatusATM.objects.create(status=f'Изменен на статус {status} (заявки)', date_change=date.today(), user=atm.user,
                                 sn=ATM.objects.get(serial_number=atm.serial_number))
    return True


def changes_req_atm_funk(req_id, sn, request):
    atm = ATM.objects.get(serial_number=sn)
    req = Request.objects.get(request_id=req_id)

    StatusATM.objects.create(
        status=f'Добавлен к заявке {req_id}',
        date_change=date.today(),
        user=request.user,
        sn=atm
    )

    ATM.objects.filter(serial_number=sn).update(
        request=req,
        status=req.status
    )

    return True


def changes_status_atm_funk(sn, new_status, request):
    try:
        # Получаем банкомат
        atm = ATM.objects.get(serial_number=sn)
        # Обновляем статус
        atm.status = new_status
        atm.save()

        # Создаём запись в истории
        StatusATM.objects.create(
            status=f'Изменен на статус {new_status} (ручное изменение)',
            date_change=date.today(),
            user=request.user,
            sn=atm
        )

        return {"success": True, "message": "Статус успешно изменён"}

    except ObjectDoesNotExist:
        return {"success": False, "message": f"Банкомат с серийным номером {sn} не найден"}

    except Exception as e:
        return {"success": False, "message": f"Ошибка при изменении статуса: {str(e)}"}


def scan_word_file(file_path, number, project, model, atm_sn, date_inv):
    doc = Document(file_path)
    d = date_inv
    formatted = f"{months[d.month]} {d.year}"

    # === ЗАМЕНЫ В АБЗАЦАХ С ШРИФТОМ ===
    # Абзац 0
    p0 = doc.paragraphs[0]
    p0.clear()
    run0 = p0.add_run(f"Акт выполненных работ по покраске банкомата № {number}")
    run0.font.name = 'Calibri'
    run0._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run0.font.size = Pt(14)
    run0.font.bold = True

    # Абзац 1
    p1 = doc.paragraphs[1]
    p1.clear()
    run1 = p1.add_run(f"{project} – Hyosung {model}  – {formatted}")
    run1.font.name = 'Calibri'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run1.font.size = Pt(14)
    run1.font.bold = True

    # Абзац 2
    p2 = doc.paragraphs[2]
    p2.clear()
    run2 = p2.add_run(f'от {d.strftime("%d.%m.%Y")}')
    run2.font.name = 'Calibri'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run2.font.size = Pt(11)
    run2.font.bold = True

    # Абзац 8
    p8 = doc.paragraphs[8]
    p8.clear()
    run8_1 = p8.add_run("На общую сумму: ")
    run8_1.font.name = 'Calibri'
    run8_1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run8_1.font.size = Pt(11)
    run8_1.font.bold = True

    total_sum = f"{len(atm_sn) * 24000:,.2f}".replace(",", " ").replace(".", ",")
    run8_2 = p8.add_run(f"{total_sum} рублей 0 копеек.")
    run8_2.font.name = 'Calibri'
    run8_2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run8_2.font.size = Pt(14)
    run8_2.font.bold = True

    # === ЗАМЕНЫ В ТАБЛИЦАХ ===
    result = ", ".join(atm_sn)
    replacements_tables = {
        (0, 1, 0): result,  # таблица 0, строка 1, столбец 0
        (1, 1, 1): f"{len(atm_sn)}",  # таблица 1, строка 1, столбец 1
        (1, 1, 2): f"24 000,00",  # таблица 1, строка 1, столбец 2
        (1, 1, 3): total_sum  # таблица 1, строка 1, столбец 3
    }

    for (t_index, r_index, c_index), new_text in replacements_tables.items():
        if t_index < len(doc.tables):
            table = doc.tables[t_index]
            if r_index < len(table.rows) and c_index < len(table.rows[r_index].cells):
                table.rows[r_index].cells[c_index].text = new_text
            else:
                print(f"Ячейка ({r_index}, {c_index}) не найдена в таблице {t_index}")
        else:
            print(f"Таблица {t_index} не найдена")

    # === СОХРАНЯЕМ ФАЙЛ ===
    doc.save(f'media/invoices/АВР_по_покраске_№{number}.docx')


def _safe_filename(name: str) -> str:
    name = (name or "UNKNOWN").strip()
    name = re.sub(r"[^\w\s\-\.]", "", name, flags=re.U)
    name = re.sub(r"\s+", "_", name).strip("_")
    return name[:80] if name else "UNKNOWN"


def _has_date(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    return bool(re.search(
        r"(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})|(\d{4}[./-]\d{1,2}[./-]\d{1,2})",
        t
    ))


def autosize_columns(ws, min_width=10, max_width=60):
    for col_idx, col_cells in enumerate(ws.columns, start=1):
        max_length = 0
        for cell in col_cells:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))

        adjusted_width = min(max_width, max(min_width, max_length + 2))
        ws.column_dimensions[get_column_letter(col_idx)].width = adjusted_width


def limit_columns_by_model(ws, header_name="Модель", default_width=30):
    model_col_idx = None

    for col_idx, cell in enumerate(ws[1], start=1):
        if str(cell.value).strip() == header_name:
            model_col_idx = col_idx
            break

    if model_col_idx:
        model_letter = get_column_letter(model_col_idx)
        max_width = ws.column_dimensions[model_letter].width or default_width
    else:
        max_width = default_width

    for col_idx in range(1, ws.max_column + 1):
        col_letter = get_column_letter(col_idx)
        ws.column_dimensions[col_letter].width = max_width

    for row in ws.iter_rows():
        for cell in row:
            if cell.value:
                cell.alignment = Alignment(wrap_text=True, vertical="top")


def format_sheet(ws, header_name="Модель", default_width=30):
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    model_col_idx = None

    for col_idx, cell in enumerate(ws[1], start=1):
        if str(cell.value).strip() == header_name:
            model_col_idx = col_idx
            break

    if model_col_idx:
        model_letter = get_column_letter(model_col_idx)
        max_width = ws.column_dimensions[model_letter].width or default_width
    else:
        max_width = default_width

    for col_idx in range(1, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = max_width

    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = border
